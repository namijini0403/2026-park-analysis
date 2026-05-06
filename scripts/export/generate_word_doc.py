# -*- coding: utf-8 -*-
"""프로젝트 구조 정리 Word 문서 생성"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# 기본 페이지 설정
section = doc.sections[0]
section.page_width = Cm(29.7)
section.page_height = Cm(21.0)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)

FONT_NAME = "맑은 고딕"

def set_font(run, size=10, bold=False, color=None):
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    try:
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn("w:eastAsia"), FONT_NAME)
    except Exception:
        pass

def add_heading(text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    sizes = {1: 16, 2: 13, 3: 11}
    colors = {1: (31, 73, 125), 2: (54, 96, 146), 3: (79, 129, 189)}
    set_font(run, size=sizes.get(level, 11), bold=True, color=colors.get(level, (0, 0, 0)))
    return p

def add_body(text, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    if indent:
        p.paragraph_format.left_indent = Cm(indent * 0.5)
    run = p.add_run(text)
    set_font(run, size=9.5)
    return p

def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)
    try:
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn("w:eastAsia"), "굴림체")
    except Exception:
        pass

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def add_table(headers, rows, header_color="1F497D", col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    # 헤더
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        shade_cell(cell, header_color)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        set_font(run, size=9, bold=True, color=(255, 255, 255))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 데이터
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg = "EBF3FB" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            shade_cell(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_font(run, size=9)
    if col_widths:
        for ri_all, row in enumerate(table.rows):
            for ci, w in enumerate(col_widths):
                if ci < len(row.cells):
                    row.cells[ci].width = Cm(w)
    doc.add_paragraph()
    return table


# ===== 제목 =====
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run("인천 초등학교 야외활동 환경 격차 분석 프로젝트\n구조 정리 문서")
set_font(title_run, size=18, bold=True, color=(31, 73, 125))

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = sub_p.add_run("작성일: 2026-04-22  |  교육공공데이터 AI활용 공모전 제출용")
set_font(sub_run, size=10, color=(128, 128, 128))

doc.add_paragraph()

# ===== 1. 프로젝트 개요 =====
add_heading("1. 프로젝트 개요", 1)
add_table(
    ["항목", "내용"],
    [
        ["목적", "교육공공데이터 AI활용 분석 공모전 제출"],
        ["주제", "인천 초등학교 주변 아동 야외활동 환경 격차 분석 및 지원 우선순위 도출"],
        ["분석 범위", "인천 전체 272개 초등학교"],
        ["핵심 지표", "도보 500m 등시선 내 공원·녹지 면적 비율 (Green Ratio)"],
        ["공간 기준", "실제 보행 도로망 기반 등시선 (Valhalla 엔진 + OSMnx)"],
        ["AI 활용", "K-means 클러스터링 / XGBoost 수혜인구 예측 / LightGBM 우선순위 보완"],
        ["팀 구성", "미추홀구 현직 초등교사(분석) + 교장(앱개발) + 발표 담당"],
        ["제출 기한", "약 45일 이내"],
    ],
    col_widths=[4, 19],
)

# ===== 2. 분석 흐름도 =====
add_heading("2. 분석 구조 흐름도", 1)
add_body("전체 분석은 6단계 파이프라인으로 구성됩니다.")

flowchart = """
┌─────────────────────────────────────────────────────────────────────┐
│                         데이터 수집 및 전처리                          │
│  공원 데이터(parks.csv) + 학교 데이터 + 인구 격자(100m) + OSM 도로망    │
└───────────────────────────┬─────────────────────────────────────────┘
                            │
                            ▼
┌─────────────────────────────────────────────────────────────────────┐
│                          공간 분석 (PHASE 2)                          │
│  OSMnx 보행 도로망 → Valhalla 등시선 생성 → 학교-공원 매핑              │
│  직선 500m 원 vs 실제 도보 등시선 비교                                  │
└───────────────────────────┬─────────────────────────────────────────┘
                            │
                            ▼
┌─────────────────────────────────────────────────────────────────────┐
│                     아파트 단지 보정 (신규)                            │
│  OSM landuse=residential → 등시선 외부 아파트 탐지                     │
│  인접 아파트 포함 → 보정 등시선(isochrone_corrected.geojson)           │
│  보정 Green Ratio 산출 → school_priority.csv 업데이트                 │
└───────────────────────────┬─────────────────────────────────────────┘
                            │
                            ▼
┌─────────────────────────────────────────────────────────────────────┐
│                      클러스터링 (PHASE 4)                             │
│  K-means k=3 (실루엣 0.4156)                                         │
│  cluster0: 접근양호형 / cluster1: 도보취약형 / cluster2: 대형공원특수형  │
│  Isolation Forest 이상치 14개 탐지                                    │
└───────────────────────────┬─────────────────────────────────────────┘
                            │
                            ▼
┌─────────────────────────────────────────────────────────────────────┐
│                      우선순위 도출 (PHASE 6)                           │
│  규칙 기반 필터 (가중치 합산 금지)                                      │
│  case_type + green_bucket + barrier_flag 조합                        │
│  priority_score → school_priority.csv                               │
└───────────────────────────┬─────────────────────────────────────────┘
                            │
                            ▼
┌─────────────────────────────────────────────────────────────────────┐
│                       앱 시각화 (index.html)                          │
│  Kakao 지도 + 보정 등시선 + 학교 마커 + 통계 패널                     │
│  app (Vite) → statisticsPreviewDataSafe.ts → 통계 카드                │
└─────────────────────────────────────────────────────────────────────┘
"""
add_code(flowchart)

# ===== 3. 데이터 파이프라인 =====
add_heading("3. 데이터 전처리 파이프라인", 1)
add_table(
    ["단계", "작업", "주요 파일/도구"],
    [
        ["PHASE 1", "파일 인코딩 정제 (CP949→UTF-8), 컬럼명 영문 통일, 결측·중복 제거", "parks.csv, 학교데이터"],
        ["PHASE 1", "좌표계 통일 (WGS84), 인천 지역 필터링, 지오코딩", "GeoPandas, Shapely"],
        ["PHASE 1", "100m 격자 인구에서 인천 영역 추출, 공원 종류 필터링", "population_grid.csv"],
        ["PHASE 2", "OSMnx 보행 도로망 수집, 등시선 생성 (Valhalla 500m)", "analysis_isochrone_valhalla.py"],
        ["PHASE 2", "학교-공원 매핑, 우회거리, barrier_flag, Voronoi 학구도", "OSMnx, NetworkX"],
        ["PHASE 2 신규", "OSM PBF → residential 폴리곤 추출 (osmium)", "incheon_residential_osm.geojson"],
        ["PHASE 2 신규", "등시선 외부 아파트 탐지 → 보정 등시선 생성", "isochrone_corrected.geojson"],
        ["PHASE 2 신규", "보정 Green Ratio 산출, school_apt_blockage.csv 생성", "analysis/fix_candidate_simulation_metrics.py"],
        ["PHASE 3", "expected_park 계산, gap 산출, gap_category 분류", "school_priority.csv"],
        ["PHASE 4", "K-means k=3 클러스터링, Isolation Forest 이상치", "sklearn"],
        ["PHASE 5", "코호트 분석, Prophet/ARIMA 인구 예측", "build_prophet_cohort_change.py"],
        ["PHASE 6", "규칙 기반 우선순위 도출, priority_score 산출", "generate_statistics_preview_data_safe.py"],
    ],
    col_widths=[3, 13, 7],
)

# ===== 4. 주요 출력 파일 =====
add_heading("4. 주요 출력 파일", 1)
add_table(
    ["파일명", "위치", "내용"],
    [
        ["school_priority.csv", "data/processed/", "학교별 전체 지표 (272개교, 50+ 컬럼)"],
        ["isochrone_corrected.geojson", "data/processed/", "아파트 보정 등시선 (224개교 확장, 48개교 미변경)"],
        ["isochrone_valhalla.geojson", "data/processed/", "원본 Valhalla 등시선"],
        ["school_apt_blockage.csv", "data/processed/", "아파트 차단 분석 결과 (교 당 보정 면적 등)"],
        ["blocked_parks_by_apt.csv", "data/processed/", "아파트에 의해 차단된 공원 목록"],
        ["incheon_residential_osm.geojson", "data/processed/", "OSM landuse=residential 폴리곤 (3,299개, 89.4km²)"],
        ["school_nearest_park.csv", "data/processed/", "학교별 최근접 공원 및 도보거리"],
        ["candidate_grid_final.ready.geojson", "data/processed/", "XGBoost 후보지 예측 격자"],
        ["gu_cohort_change_prophet.csv", "data/processed/", "구별 아동인구 코호트 변화 (Prophet)"],
        ["statisticsPreviewDataSafe.ts", "app/src/", "앱 통계 패널 데이터 (빌드 결과)"],
    ],
    col_widths=[7, 5, 11],
)

# ===== 5. case_type 분류 =====
add_heading("5. case_type 분류 체계", 1)
add_body("녹지 도보접근성 기준 4단계 분류 (green_ratio = 보정 등시선 내 공원면적 / 등시선면적 × 100)")
add_table(
    ["case_type", "label", "green_ratio 기준", "해석"],
    [
        ["case1", "공원접근결핍", "최근접 공원 500m 이상 + 등시선 내 공원 0개 + 녹지 0%", "실질적 공원 접근 결핍"],
        ["case2", "녹지부족", "green_ratio < 1%", "녹지 극히 부족"],
        ["case3", "중간", "1% ≤ green_ratio < 5%", "일부 녹지 있으나 기준 미달"],
        ["case4", "양호", "green_ratio ≥ 5%", "상대적으로 녹지 충분"],
    ],
    col_widths=[3, 4, 6, 10],
)

add_heading("5-1. 클러스터 명칭 (K-means k=3)", 2)
add_table(
    ["클러스터", "명칭", "특징"],
    [
        ["cluster 0", "접근양호형", "녹지·공원 접근성 양호, 도보거리 짧음"],
        ["cluster 1", "도보취약형", "등시선 내 녹지 부족, 실제 도보거리 길거나 장벽 있음"],
        ["cluster 2", "대형공원특수형", "대형 공원 인접하나 접근 경로 특수 (이상치 포함)"],
    ],
    col_widths=[3, 5, 15],
)

# ===== 6. 아파트 보정 방법론 =====
add_heading("6. 아파트 단지 보정 방법론", 1)
add_body("인천 아파트 단지는 보행이 자유로우므로, 단지 내부를 통과하는 도보권을 보정하여 실질 등시선을 확장합니다.")

add_heading("6-1. 보정 공식", 2)
add_code("gap = 500m 직선 버퍼 − 원본 등시선")
add_code("apt_in_gap = gap ∩ residential_union  (등시선 외부의 아파트 면적)")
add_code("corrected_isochrone = original_iso ∪ apt_in_gap_touching")
add_code("  (touching = iso.buffer(2m)와 교차하는 아파트 폴리곤만 포함 → 연결성 보장)")
add_code("")
add_code("corrected_green_ratio = (iso_park_area + blocked_intersect_m2)")
add_code("                        / min(isochrone_area_m2 + apt_blocked_m2, 785398) × 100")
add_code("  (785398 = π × 500² = 직선 500m 원 최대면적)")

add_heading("6-2. 보정 제외 대상 (보호 학교)", 2)
add_table(
    ["학교ID", "학교명", "제외 사유"],
    [
        ["B000003144", "인천석남초등학교", "실측 보정값 적용 중"],
        ["B000003145", "인천천마초등학교", "실측 보정값 적용 중"],
        ["B000026504", "인천송담초등학교", "실측 보정값 적용 중"],
        ["B000003048", "인천명선초등학교", "실측 보정값 적용 중"],
        ["B000002990", "인천학익초등학교", "실측 보정값 적용 중"],
    ],
    col_widths=[4, 6, 13],
)

# ===== 7. 앱 구조 =====
add_heading("7. 앱 구조", 1)

add_heading("7-1. 메인 앱 (index.html)", 2)
add_table(
    ["구성 요소", "역할"],
    [
        ["Kakao 지도", "학교 마커, 등시선(보정), 500m 원, 공원 레이어 시각화"],
        ["통계 패널", "학교 클릭 시 상세 정보 (녹지율, 최근접 공원, case_type, 우선순위 등)"],
        ["필터 UI", "구 선택, case_type 필터, 클러스터 필터"],
        ["공유 URL", "지도 상태(학교ID, 중심좌표, 줌) 해시 파라미터로 공유 가능"],
        ["데이터 로딩", "isochrone_corrected.geojson + parks.csv + school_priority.csv"],
    ],
    col_widths=[6, 17],
)

add_heading("7-2. React 상세 리포트 앱 (app)", 2)
add_table(
    ["파일", "역할"],
    [
        ["app/src/statisticsPreviewDataSafe.ts", "전체 통계 데이터 (빌드된 결과 → iframe 앱에 반영)"],
        ["scripts/export/generate_statistics_preview_data_safe.py", "통계 TS 파일 생성 스크립트 (school_priority.csv → TS)"],
        ["app/vite.config.ts", "Vite 빌드 설정"],
    ],
    col_widths=[9, 14],
)

add_heading("7-3. 통계 업데이트 절차", 2)
add_body("school_priority.csv 변경 후 아래 순서로 실행:")
add_code("1. python scripts/export/generate_statistics_preview_data_safe.py")
add_code("2. cd app && npm run build")
add_code("3. git add -A && git commit && git push")
add_body("→ GitHub Pages에 자동 반영 (index.html 내 statisticsPreviewDataSafe.ts 임베드)")

# ===== 8. 핵심 지표 정의 =====
add_heading("8. 핵심 지표 정의", 1)
add_table(
    ["지표명", "정의", "비고"],
    [
        ["iso_green_ratio", "보정 등시선 내 공원면적 / 등시선면적 × 100 (%)", "메인 지표"],
        ["iso_park_area", "보정 등시선 내 공원 면적 합계 (m²)", ""],
        ["iso_park_count", "보정 등시선 내 공원 수", ""],
        ["nearest_park_dist_m", "학교에서 가장 가까운 공원까지 도보거리 (m)", "실측 봉인값 우선"],
        ["apt_blocked_m2", "등시선 외부 아파트 단지가 차단하는 면적 (m²)", "보정 신규"],
        ["corrected_green_ratio", "아파트 보정 후 green ratio (%)", "보정 신규"],
        ["priority_score", "규칙 기반 우선순위 점수 (낮을수록 우선)", ""],
        ["case_type", "녹지접근성 4단계 분류 (case1~4)", ""],
        ["cluster", "K-means 클러스터 (0=양호, 1=취약, 2=특수)", ""],
        ["barrier_flag", "도보 경로 장벽 존재 여부 (True/False)", ""],
    ],
    col_widths=[5, 12, 6],
)

# ===== 9. 봉인 실측값 =====
add_heading("9. 봉인된 실측값 (변경 금지)", 1)
add_body("아래 값은 사용자가 네이버지도 기준으로 직접 실측한 값입니다. 파이프라인 재실행 시에도 덮어쓰지 말 것.")
add_table(
    ["학교명", "실제 최근접 공원", "실측 도보거리(m)", "기존 오류값(m)", "확인일"],
    [
        ["인천신석초등학교", "석남도시숲", "120", "664.8", "2026-04-16"],
        ["인천석남초등학교", "석곶체육공원", "444", "1010.3", "2026-04-16"],
        ["인천새봄초등학교", "동춘1구역근린공원", "50", "1812.1", "2026-04-18"],
        ["인천만석초등학교", "화도진공원", "63", "110.6", "2026-04-18"],
        ["인천경원초등학교", "다솔어린이공원", "88", "73.2", "2026-04-18"],
    ],
    col_widths=[5, 6, 4, 4, 4],
)

# ===== 10. XGBoost 후보지 모델 =====
add_heading("10. XGBoost 후보지 모델 (candidate_grid)", 1)
add_body("공원 부족 지역 내 신규 공원 후보지 격자(100m×100m)를 XGBoost로 예측합니다.")
add_table(
    ["항목", "내용"],
    [
        ["입력", "candidate_grid_final.ready.geojson (격자별 특성값)"],
        ["피처", "격자 총인구, 재개발 상태, 구 더미, 등시선 내 시설 수, 도로망 밀도"],
        ["타겟", "학교별 등시선 내 0~12세 아동인구 (수혜인구)"],
        ["출력", "data/processed/candidate_grid_xgb_v4.tmp.geojson"],
        ["시각화", "index.html 지도 레이어 (후보지 격자 열지도)"],
        ["SHAP", "피처 중요도 시각화 (분석 보고서용)"],
    ],
    col_widths=[4, 19],
)

# 저장
output_path = r"c:\2026_data_analysis_park\프로젝트_구조_정리_20260422.docx"
doc.save(output_path)
print(f"저장 완료: {output_path}")
